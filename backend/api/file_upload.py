"""
File Upload and Document Processing API for CEP Machine
Production-ready file handling with text extraction and processing
"""

from fastapi import APIRouter, UploadFile, File, HTTPException, BackgroundTasks
from fastapi.responses import JSONResponse, FileResponse
import aiofiles
import os
from typing import List, Optional
from datetime import datetime
import uuid
import hashlib
import logging
import mimetypes

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/files", tags=["files"])

# Configuration
UPLOAD_DIR = os.getenv("UPLOAD_DIR", "/app/uploads")
MAX_FILE_SIZE = int(os.getenv("MAX_FILE_SIZE_MB", "10")) * 1024 * 1024  # Default 10MB
ALLOWED_EXTENSIONS = {
    ".pdf", ".docx", ".doc", ".txt", ".csv", ".json", 
    ".xlsx", ".xls", ".md", ".rtf"
}
ALLOWED_MIME_TYPES = {
    "application/pdf",
    "application/msword",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
    "text/csv",
    "application/json",
    "application/vnd.ms-excel",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "text/markdown"
}

def get_file_hash(content: bytes) -> str:
    """Generate SHA-256 hash of file content"""
    return hashlib.sha256(content).hexdigest()

def validate_file(filename: str, content: bytes) -> tuple[bool, str]:
    """Validate file type and size"""
    # Check file size
    if len(content) > MAX_FILE_SIZE:
        return False, f"File size exceeds maximum allowed ({MAX_FILE_SIZE // (1024*1024)}MB)"
    
    # Check extension
    file_extension = os.path.splitext(filename)[1].lower()
    if file_extension not in ALLOWED_EXTENSIONS:
        return False, f"File type '{file_extension}' not allowed. Allowed types: {', '.join(ALLOWED_EXTENSIONS)}"
    
    return True, "Valid"

async def extract_text_from_file(file_path: str) -> str:
    """Extract text from uploaded file"""
    file_extension = os.path.splitext(file_path)[1].lower()
    
    try:
        if file_extension == ".pdf":
            return await _extract_pdf_text(file_path)
        elif file_extension in [".docx", ".doc"]:
            return await _extract_docx_text(file_path)
        elif file_extension == ".txt":
            return await _extract_txt_text(file_path)
        elif file_extension == ".csv":
            return await _extract_csv_text(file_path)
        elif file_extension == ".json":
            return await _extract_json_text(file_path)
        elif file_extension in [".xlsx", ".xls"]:
            return await _extract_excel_text(file_path)
        elif file_extension == ".md":
            return await _extract_txt_text(file_path)
        else:
            return f"Text extraction not supported for {file_extension} files"
    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {str(e)}")
        return f"Error extracting text: {str(e)}"

async def _extract_pdf_text(file_path: str) -> str:
    """Extract text from PDF file"""
    try:
        import PyPDF2
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text += f"\n--- Page {page_num + 1} ---\n{page_text}"
        return text.strip() or "No text content found in PDF"
    except ImportError:
        return "PyPDF2 not installed. Cannot extract PDF text."
    except Exception as e:
        return f"Error reading PDF: {str(e)}"

async def _extract_docx_text(file_path: str) -> str:
    """Extract text from DOCX file"""
    try:
        from docx import Document
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                text += " | ".join(row_text) + "\n"
        
        return text.strip() or "No text content found in document"
    except ImportError:
        return "python-docx not installed. Cannot extract DOCX text."
    except Exception as e:
        return f"Error reading DOCX: {str(e)}"

async def _extract_txt_text(file_path: str) -> str:
    """Extract text from TXT file"""
    try:
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
            return await file.read()
    except UnicodeDecodeError:
        # Try with different encoding
        async with aiofiles.open(file_path, 'r', encoding='latin-1') as file:
            return await file.read()

async def _extract_csv_text(file_path: str) -> str:
    """Extract text from CSV file"""
    try:
        import pandas as pd
        df = pd.read_csv(file_path)
        return f"CSV with {len(df)} rows and {len(df.columns)} columns:\n\nColumns: {', '.join(df.columns)}\n\nPreview:\n{df.head(10).to_string()}"
    except ImportError:
        # Fallback without pandas
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
            return await file.read()
    except Exception as e:
        return f"Error reading CSV: {str(e)}"

async def _extract_json_text(file_path: str) -> str:
    """Extract text from JSON file"""
    import json
    try:
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
            content = await file.read()
            data = json.loads(content)
            return json.dumps(data, indent=2)
    except Exception as e:
        return f"Error reading JSON: {str(e)}"

async def _extract_excel_text(file_path: str) -> str:
    """Extract text from Excel file"""
    try:
        import pandas as pd
        # Read all sheets
        xlsx = pd.ExcelFile(file_path)
        text_parts = []
        
        for sheet_name in xlsx.sheet_names:
            df = pd.read_excel(xlsx, sheet_name=sheet_name)
            text_parts.append(f"\n--- Sheet: {sheet_name} ---\n")
            text_parts.append(f"Rows: {len(df)}, Columns: {len(df.columns)}\n")
            text_parts.append(f"Columns: {', '.join(df.columns.astype(str))}\n")
            text_parts.append(f"Preview:\n{df.head(10).to_string()}\n")
        
        return "".join(text_parts)
    except ImportError:
        return "pandas/openpyxl not installed. Cannot extract Excel text."
    except Exception as e:
        return f"Error reading Excel: {str(e)}"

async def save_upload_file(upload_file: UploadFile) -> dict:
    """Save uploaded file and return metadata"""
    # Read file content
    content = await upload_file.read()
    
    # Validate file
    is_valid, message = validate_file(upload_file.filename, content)
    if not is_valid:
        raise HTTPException(status_code=400, detail=message)
    
    # Generate unique filename
    file_extension = os.path.splitext(upload_file.filename)[1].lower()
    file_id = str(uuid.uuid4())
    unique_filename = f"{file_id}{file_extension}"
    
    # Create upload directory if it doesn't exist
    os.makedirs(UPLOAD_DIR, exist_ok=True)
    
    # Save file
    file_path = os.path.join(UPLOAD_DIR, unique_filename)
    async with aiofiles.open(file_path, 'wb') as f:
        await f.write(content)
    
    # Generate metadata
    file_hash = get_file_hash(content)
    
    return {
        "file_id": file_id,
        "original_filename": upload_file.filename,
        "stored_filename": unique_filename,
        "file_path": file_path,
        "file_size": len(content),
        "file_hash": file_hash,
        "content_type": upload_file.content_type,
        "extension": file_extension,
        "uploaded_at": datetime.utcnow().isoformat()
    }

@router.post("/upload")
async def upload_file(
    file: UploadFile = File(...),
    extract_text: bool = True
):
    """Upload and optionally process a file"""
    try:
        # Save file
        file_metadata = await save_upload_file(file)
        
        response = {
            "success": True,
            "file_id": file_metadata["file_id"],
            "filename": file_metadata["original_filename"],
            "size": file_metadata["file_size"],
            "content_type": file_metadata["content_type"],
            "uploaded_at": file_metadata["uploaded_at"],
            "message": "File uploaded successfully"
        }
        
        # Extract text if requested
        if extract_text:
            extracted_text = await extract_text_from_file(file_metadata["file_path"])
            response["extracted_text"] = extracted_text[:5000]  # Limit to first 5000 chars
            response["text_length"] = len(extracted_text)
            response["text_truncated"] = len(extracted_text) > 5000
        
        logger.info(f"File uploaded: {file_metadata['original_filename']} ({file_metadata['file_size']} bytes)")
        return JSONResponse(response)
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error uploading file: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error uploading file: {str(e)}")

@router.post("/upload/multiple")
async def upload_multiple_files(
    files: List[UploadFile] = File(...),
    extract_text: bool = False
):
    """Upload multiple files"""
    results = []
    errors = []
    
    for file in files:
        try:
            file_metadata = await save_upload_file(file)
            
            result = {
                "success": True,
                "file_id": file_metadata["file_id"],
                "filename": file_metadata["original_filename"],
                "size": file_metadata["file_size"]
            }
            
            if extract_text:
                extracted_text = await extract_text_from_file(file_metadata["file_path"])
                result["extracted_text"] = extracted_text[:2000]
                result["text_truncated"] = len(extracted_text) > 2000
            
            results.append(result)
            
        except Exception as e:
            errors.append({
                "filename": file.filename,
                "error": str(e)
            })
    
    return JSONResponse({
        "success": len(errors) == 0,
        "uploaded": len(results),
        "failed": len(errors),
        "files": results,
        "errors": errors
    })

@router.get("/list")
async def list_files(
    limit: int = 50,
    offset: int = 0
):
    """List uploaded files"""
    try:
        files = []
        
        if os.path.exists(UPLOAD_DIR):
            all_files = sorted(
                os.listdir(UPLOAD_DIR),
                key=lambda x: os.path.getmtime(os.path.join(UPLOAD_DIR, x)),
                reverse=True
            )
            
            for filename in all_files[offset:offset + limit]:
                file_path = os.path.join(UPLOAD_DIR, filename)
                if os.path.isfile(file_path):
                    stat = os.stat(file_path)
                    files.append({
                        "filename": filename,
                        "size": stat.st_size,
                        "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                        "created": datetime.fromtimestamp(stat.st_ctime).isoformat()
                    })
        
        return {
            "success": True,
            "files": files,
            "total": len(os.listdir(UPLOAD_DIR)) if os.path.exists(UPLOAD_DIR) else 0,
            "limit": limit,
            "offset": offset
        }
    
    except Exception as e:
        logger.error(f"Error listing files: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error listing files: {str(e)}")

@router.get("/{file_id}")
async def get_file_info(file_id: str):
    """Get information about a specific file"""
    try:
        if not os.path.exists(UPLOAD_DIR):
            raise HTTPException(status_code=404, detail="File not found")
        
        # Find file by ID
        for filename in os.listdir(UPLOAD_DIR):
            if filename.startswith(file_id):
                file_path = os.path.join(UPLOAD_DIR, filename)
                stat = os.stat(file_path)
                
                return {
                    "success": True,
                    "file_id": file_id,
                    "filename": filename,
                    "size": stat.st_size,
                    "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                    "extension": os.path.splitext(filename)[1]
                }
        
        raise HTTPException(status_code=404, detail="File not found")
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting file info: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error getting file info: {str(e)}")

@router.get("/{file_id}/download")
async def download_file(file_id: str):
    """Download a file by ID"""
    try:
        if not os.path.exists(UPLOAD_DIR):
            raise HTTPException(status_code=404, detail="File not found")
        
        # Find file by ID
        for filename in os.listdir(UPLOAD_DIR):
            if filename.startswith(file_id):
                file_path = os.path.join(UPLOAD_DIR, filename)
                return FileResponse(
                    file_path,
                    filename=filename,
                    media_type=mimetypes.guess_type(filename)[0] or "application/octet-stream"
                )
        
        raise HTTPException(status_code=404, detail="File not found")
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error downloading file: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error downloading file: {str(e)}")

@router.get("/{file_id}/extract")
async def extract_file_text(file_id: str):
    """Extract text from a file"""
    try:
        if not os.path.exists(UPLOAD_DIR):
            raise HTTPException(status_code=404, detail="File not found")
        
        # Find file by ID
        for filename in os.listdir(UPLOAD_DIR):
            if filename.startswith(file_id):
                file_path = os.path.join(UPLOAD_DIR, filename)
                extracted_text = await extract_text_from_file(file_path)
                
                return {
                    "success": True,
                    "file_id": file_id,
                    "filename": filename,
                    "text": extracted_text,
                    "text_length": len(extracted_text)
                }
        
        raise HTTPException(status_code=404, detail="File not found")
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error extracting text: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error extracting text: {str(e)}")

@router.delete("/{file_id}")
async def delete_file(file_id: str):
    """Delete a file by ID"""
    try:
        if not os.path.exists(UPLOAD_DIR):
            raise HTTPException(status_code=404, detail="File not found")
        
        # Find and delete file by ID
        for filename in os.listdir(UPLOAD_DIR):
            if filename.startswith(file_id):
                file_path = os.path.join(UPLOAD_DIR, filename)
                os.remove(file_path)
                logger.info(f"File deleted: {filename}")
                
                return {
                    "success": True,
                    "message": "File deleted successfully",
                    "file_id": file_id
                }
        
        raise HTTPException(status_code=404, detail="File not found")
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting file: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error deleting file: {str(e)}")
